import io
import yaml
import warnings
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def add_hyperlink(paragraph, url, text, color="0000FF", underline=True, font_name="Calibri", font_size=Pt(14)):
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rPr.append(c)

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size.pt * 2)))
    rPr.append(sz)

    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph.add_run("")


def add_horizontal_line(doc):
    paragraph = doc.add_paragraph()
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_left_right_paragraph(doc, left_text, right_text,
                             left_bold=False, right_bold=False,
                             left_italic=False, right_italic=False,
                             font_size=12, right_color=None):
    paragraph = doc.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.tab_stops.clear_all()
    paragraph_format.tab_stops.add_tab_stop(Inches(6.5), alignment=WD_TAB_ALIGNMENT.RIGHT)

    left_run = paragraph.add_run(left_text)
    left_run.font.name = "Calibri"
    left_run.font.size = Pt(font_size)
    left_run.bold = left_bold
    left_run.italic = left_italic

    paragraph.add_run("\t")

    right_run = paragraph.add_run(right_text)
    right_run.font.name = "Calibri"
    right_run.font.size = Pt(font_size)
    right_run.bold = right_bold
    right_run.italic = right_italic

    if right_color:
        right_run.font.color.rgb = right_color

    return paragraph


def create_resume_from_yaml(yaml_data, doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1

    if 'List Bullet' in doc.styles:
        bullet_style = doc.styles['List Bullet']
        bullet_style.font.name = 'Calibri'
        bullet_style.font.size = Pt(12)
        bullet_style.paragraph_format.line_spacing = 1
        bullet_style.paragraph_format.space_after = Pt(0)

    # 1) HEADER
    header_data = yaml_data["resumeStructure"].get("header", {})
    name_line = header_data.get("line1", "")
    contact_line = header_data.get("line2", "")

    name_paragraph = doc.add_paragraph()
    name_run = name_paragraph.add_run(name_line)
    name_run.font.name = "Calibri"
    name_run.font.size = Pt(28)
    name_run.bold = True
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact_paragraph = doc.add_paragraph()
    parts = [p.strip() for p in contact_line.split(" | ")]

    for i, part in enumerate(parts):
        if i > 0:
            sep_run = contact_paragraph.add_run(" | ")
            sep_run.font.name = "Calibri"
            sep_run.font.size = Pt(14)

        if part.startswith("http"):
            # Derive display text from URL
            lower = part.lower()
            if "linkedin" in lower:
                display = "LinkedIn"
            elif "github" in lower:
                display = "GitHub"
            else:
                display = "Portfolio"
            link_run = add_hyperlink(contact_paragraph, part, display)
            link_run.font.name = "Calibri"
            link_run.font.size = Pt(14)
        else:
            text_run = contact_paragraph.add_run(part)
            text_run.font.name = "Calibri"
            text_run.font.size = Pt(14)

    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_horizontal_line(doc)

    # 2) WORK EXPERIENCE
    work_experience_data = yaml_data["resumeStructure"].get("workExperience", {})
    work_experience_header = work_experience_data.get("header", "WORK EXPERIENCE")

    we_header_para = doc.add_paragraph()
    we_header_para.paragraph_format.space_after = Pt(0)
    we_run = we_header_para.add_run(work_experience_header.upper())
    we_run.font.name = "Calibri"
    we_run.font.size = Pt(12)
    we_run.bold = True
    we_header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_horizontal_line(doc)

    companies = work_experience_data.get("companies", [])
    for c_idx, company in enumerate(companies):
        company_info_line = company.get("companyInfoLine", "")
        parts = [p.strip() for p in company_info_line.split("|")]

        if len(parts) == 2:
            dates_worked, company_name = parts[0], parts[1]
        else:
            dates_worked = parts[0] if parts else ""
            company_name = parts[1] if len(parts) > 1 else ""

        add_left_right_paragraph(
            doc, left_text=dates_worked, right_text=company_name,
            left_bold=True, right_bold=True, font_size=12
        )

        roles = company.get("roles", [])
        for r_idx, role in enumerate(roles):
            role_info_line = role.get("roleInfoLine", "")
            parts = [p.strip() for p in role_info_line.split("|")]

            role_name = parts[0] if len(parts) > 0 else ""
            role_dates = parts[1] if len(parts) > 1 else ""
            role_city_state = parts[2] if len(parts) > 2 else ""

            add_left_right_paragraph(
                doc, left_text=f"{role_name} | {role_dates}",
                right_text=role_city_state,
                left_italic=True, right_italic=True, font_size=12
            )

            bullets = role.get("bullets", [])
            for bullet_item in bullets:
                bullet_description = bullet_item.get("description", "")
                bullet_para = doc.add_paragraph(style='List Bullet')
                bullet_para.paragraph_format.space_after = Pt(0)
                bullet_run = bullet_para.add_run(bullet_description)
                bullet_run.font.name = "Calibri"
                bullet_run.font.size = Pt(12)

                sub_bullets = bullet_item.get("subBullets", {})
                if isinstance(sub_bullets, dict):
                    sub_bullets = [sub_bullets]
                elif not isinstance(sub_bullets, list):
                    sub_bullets = []

                for sub_bullet in sub_bullets:
                    if isinstance(sub_bullet, dict):
                        for key, value in sub_bullet.items():
                            if key == "techStack":
                                sub_bullet_para = doc.add_paragraph(style='List Bullet 2')
                                sub_bullet_para.paragraph_format.space_after = Pt(0)
                                label_run = sub_bullet_para.add_run("Tech Stack: ")
                                label_run.bold = True
                                label_run.font.name = "Calibri"
                                value_run = sub_bullet_para.add_run(value)
                                value_run.font.name = "Calibri"
                            elif key == "keyResult":
                                sub_bullet_para = doc.add_paragraph(style='List Bullet 2')
                                sub_bullet_para.paragraph_format.space_after = Pt(0)
                                label_run = sub_bullet_para.add_run("Key Results: ")
                                label_run.bold = True
                                label_run.font.name = "Calibri"
                                value_run = sub_bullet_para.add_run(value)
                                value_run.font.name = "Calibri"

            if r_idx < len(roles) - 1:
                doc.add_paragraph()

        if c_idx < len(companies) - 1:
            doc.add_paragraph()

    doc.add_paragraph()

    # 3) EDUCATION
    education_data = yaml_data["resumeStructure"].get("education", {})
    education_header = education_data.get("header", "EDUCATION")

    ed_header_para = doc.add_paragraph()
    ed_header_para.paragraph_format.space_after = Pt(0)
    ed_run = ed_header_para.add_run(education_header.upper())
    ed_run.font.name = "Calibri"
    ed_run.font.size = Pt(12)
    ed_run.bold = True
    ed_header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_horizontal_line(doc)

    university_name_line = education_data.get("universityNameLine", "")
    degree_info_line = education_data.get("degreeInfoLine", "")

    univ_para = doc.add_paragraph()
    univ_run = univ_para.add_run(university_name_line)
    univ_run.font.name = "Calibri"
    univ_run.font.size = Pt(12)
    univ_run.bold = True

    parts = [p.strip() for p in degree_info_line.split("|")]
    deg = parts[0] if len(parts) > 0 else ""
    location = parts[1] if len(parts) > 1 else ""

    add_left_right_paragraph(
        doc, left_text=deg, right_text=location,
        left_italic=True, right_italic=True, font_size=12
    )

    doc.add_paragraph()

    # 4) CERTIFICATIONS & SKILLS
    csi_data = yaml_data["resumeStructure"].get("certificationsSkills", {})
    csi_header = csi_data.get("header", "CERTIFICATIONS & SKILLS")

    csi_header_para = doc.add_paragraph()
    csi_header_para.paragraph_format.space_after = Pt(0)
    csi_run = csi_header_para.add_run(csi_header.upper())
    csi_run.font.name = "Calibri"
    csi_run.font.size = Pt(12)
    csi_run.bold = True
    csi_header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_horizontal_line(doc)

    csi_bullets = csi_data.get("bullets", [])
    for bullet_obj in csi_bullets:
        for key, val in bullet_obj.items():
            if key.lower() != "interests" and val.strip():
                para = doc.add_paragraph(style='List Bullet')
                label_run = para.add_run(f"{key.title()}: ")
                label_run.font.name = "Calibri"
                label_run.font.size = Pt(12)
                label_run.bold = True
                info_run = para.add_run(val)
                info_run.font.name = "Calibri"
                info_run.font.size = Pt(12)


def convert_yaml_to_word_bytes(yaml_content: str, company_name: str) -> tuple[bytes, str]:
    if not company_name:
        raise ValueError("Company name cannot be empty!")

    yaml_data = yaml.safe_load(yaml_content)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.gutter = 0

    create_resume_from_yaml(yaml_data, doc)

    # Extract name from YAML for filename
    name = "Resume"
    try:
        name = yaml_data["resumeStructure"]["header"]["line1"]
    except (KeyError, TypeError):
        pass

    today_date = datetime.now().strftime("%Y-%m-%d")
    filename = f"{name} Resume - {company_name} - {today_date}.docx"

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue(), filename
