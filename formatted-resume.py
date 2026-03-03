import os
import yaml  # ...added import...
import warnings
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True, font_name="Calibri", font_size=Pt(14)):
    """
    A safer hyperlink creation function for python-docx that reduces the chance of 
    producing a corrupted document. 
    - paragraph: the paragraph we want to add the hyperlink into.
    - url: the external hyperlink address.
    - text: the text displayed in the document for the hyperlink.
    - color: hex color code for the link text. Default is "0000FF" (blue).
    - underline: whether the hyperlink text is underlined. Default is True.
    - font_name: the font name to apply to the hyperlink text. Default is "Calibri".
    - font_size: the font size to apply to the hyperlink text. Default is 16pt.

    Returns a Run object, allowing additional styling (font name, size, bold, etc.).
    """
    # Create the relationship for the URL
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and set the relationship ID
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r (run) element
    new_run = OxmlElement('w:r')
    # Create a w:rPr (run properties) element
    rPr = OxmlElement('w:rPr')

    # Underline if desired
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    # Set the link color
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rPr.append(c)

    # Apply font name and size
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size.pt * 2)))  # Word stores font size in half-points
    rPr.append(sz)

    new_run.append(rPr)

    # Create a w:t element containing the display text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    # Append this run to the hyperlink
    hyperlink.append(new_run)

    # Add the hyperlink into the paragraph
    paragraph._p.append(hyperlink)

    # Return an empty Run object for spacing/styling, if desired
    return paragraph.add_run("")


def add_horizontal_line(doc):
    """
    Inserts a full-width horizontal line by applying a bottom border to a paragraph.
    """
    paragraph = doc.add_paragraph()
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')    # border thickness
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pbdr.append(bottom)
    pPr.append(pbdr)



def add_left_right_paragraph(doc, left_text, right_text, 
                             left_bold=False, right_bold=False,
                             left_italic=False, right_italic=False,
                             font_size=12, right_color=None):
    """
    Creates a paragraph with a tab stop so that the text before the tab is left-aligned
    and the text after the tab is right-aligned. This avoids using tables for ATS compatibility.
    """
    paragraph = doc.add_paragraph()
    paragraph_format = paragraph.paragraph_format

    # Clear existing tab stops, then add one near the right margin (6.5" is typical for an 8.5" wide page minus margins).
    paragraph_format.tab_stops.clear_all()
    paragraph_format.tab_stops.add_tab_stop(Inches(6.5), alignment=WD_TAB_ALIGNMENT.RIGHT)

    # Left run
    left_run = paragraph.add_run(left_text)
    left_run.font.name = "Calibri"
    left_run.font.size = Pt(font_size)
    left_run.bold = left_bold
    left_run.italic = left_italic

    # Insert a tab character to move subsequent text to the right tab stop
    paragraph.add_run("\t")

    # Right run
    right_run = paragraph.add_run(right_text)
    right_run.font.name = "Calibri"
    right_run.font.size = Pt(font_size)
    right_run.bold = right_bold
    right_run.italic = right_italic

    if right_color:
        right_run.font.color.rgb = right_color

    return paragraph

def create_resume_from_yaml(yaml_data, doc):
    """
    Creates a styled Word resume from the provided YAML data (no tables),
    using tab stops to simulate left/right alignment. 
    """

    # ========== Global Paragraph/Bullet Style Adjustments ==========
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1

    if 'List Bullet' in doc.styles:
        bullet_style = doc.styles['List Bullet']
        bullet_style.font.name = 'Calibri'
        bullet_style.font.size = Pt(12)
        bullet_style.paragraph_format.line_spacing = 1
        bullet_style.paragraph_format.space_after = Pt(0)

    # ========== 1) HEADER (Name + Contact) ==========
    header_data = yaml_data["resumeStructure"].get("header", {})
    name_line = header_data.get("line1", "")
    contact_line = header_data.get("line2", "")

    # Name: 28pt, bold
    name_paragraph = doc.add_paragraph()
    name_run = name_paragraph.add_run(name_line)
    name_run.font.name = "Calibri"
    name_run.font.size = Pt(28)
    name_run.bold = True
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Contact info: 14pt
    contact_paragraph = doc.add_paragraph()
    parts = contact_line.split(" | ")
    normal_text_parts = parts[:-1]  # everything except last
    portfolio_url = parts[-1] if parts else ""

    # Add the normal text (e.g., email, phone, city/state)
    contact_run = contact_paragraph.add_run(" | ".join(normal_text_parts))
    contact_run.font.name = "Calibri"
    contact_run.font.size = Pt(14)

    # If last part is a valid URL, create a hyperlink
    if portfolio_url.startswith("http"):
        contact_paragraph.add_run(" | ")
        link_run = add_hyperlink(contact_paragraph, portfolio_url, "Portfolio Link")
        # Style the returned run so it matches 14pt Garamond
        link_run.font.name = "Calibri"
        link_run.font.size = Pt(14)
    else:
        # Otherwise, just treat it as normal text
        if portfolio_url:
            run_portfolio = contact_paragraph.add_run(" | " + portfolio_url)
            run_portfolio.font.name = "Calibri"
            run_portfolio.font.size = Pt(14)

    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Horizontal line after header
    add_horizontal_line(doc)

    # ========== 2) WORK EXPERIENCE ==========
    work_experience_data = yaml_data["resumeStructure"].get("workExperience", {})
    work_experience_header = work_experience_data.get("header", "WORK EXPERIENCE")

    we_header_para = doc.add_paragraph()
    we_header_para.paragraph_format.space_after = Pt(0)  # Reduce space after the header
    we_run = we_header_para.add_run(work_experience_header.upper())
    we_run.font.name = "Calibri"
    we_run.font.size = Pt(12)
    we_run.bold = True
    we_header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_horizontal_line(doc)

    companies = work_experience_data.get("companies", [])
    for c_idx, company in enumerate(companies):
        company_info_line = company.get("companyInfoLine", "")
        parts = [p.strip() for p in company_info_line.split("|")]

        # "Dates | Company Name" or similar
        if len(parts) == 2:
            dates_worked, company_name = parts[0], parts[1]
        else:
            dates_worked = parts[0] if parts else ""
            company_name = parts[1] if len(parts) > 1 else ""

        # Left = dates, Right = company
        add_left_right_paragraph(
            doc,
            left_text=dates_worked,
            right_text=company_name,
            left_bold=True,
            right_bold=True,
            font_size=12
        )

        roles = company.get("roles", [])
        for r_idx, role in enumerate(roles):
            role_info_line = role.get("roleInfoLine", "")
            parts = [p.strip() for p in role_info_line.split("|")]

            role_name = parts[0] if len(parts) > 0 else ""
            role_dates = parts[1] if len(parts) > 1 else ""
            role_city_state = parts[2] if len(parts) > 2 else ""

            # Left = Role Name + Role Dates (italic), Right = City/State (italic)
            add_left_right_paragraph(
                doc,
                left_text=f"{role_name} | {role_dates}",
                right_text=role_city_state,
                left_italic=True,
                right_italic=True,
                font_size=12
            )

            # Bullets
            bullets = role.get("bullets", [])
            for bullet_item in bullets:
                bullet_description = bullet_item.get("description", "")
                bullet_para = doc.add_paragraph(style='List Bullet')
                bullet_para.paragraph_format.space_after = Pt(0)  # Set 0pt spacing after each bullet
                bullet_run = bullet_para.add_run(bullet_description)
                bullet_run.font.name = "Calibri"
                bullet_run.font.size = Pt(12)

                sub_bullets = bullet_item.get("subBullets", {})
                # Ensure sub_bullets is always iterable (convert dictionary to list)
                if isinstance(sub_bullets, dict):
                    sub_bullets = [sub_bullets]  # Wrap dictionary in a list for uniform handling

                elif not isinstance(sub_bullets, list):
                    sub_bullets = []  # Default to empty list if unexpected format



                # Process sub-bullets
                for sub_bullet in sub_bullets:
                    if isinstance(sub_bullet, dict):  # Ensure each item is a dictionary
                        for key, value in sub_bullet.items():
                         if key == "techStack":
                            sub_bullet_para = doc.add_paragraph(style='List Bullet 2')
                            sub_bullet_para.paragraph_format.space_after = Pt(0)
                            label_run = sub_bullet_para.add_run("Tech Stack: ")
                            label_run.bold = True
                            label_run.font.name = "Calibri"
                            value_run = sub_bullet_para.add_run(value)
                            value_run.font.name = "Calibri"
                         elif key == "keyResult":
                            sub_bullet_para = doc.add_paragraph(style='List Bullet 2')
                            sub_bullet_para.paragraph_format.space_after = Pt(0)
                            label_run = sub_bullet_para.add_run("Key Results: ")
                            label_run.bold = True
                            label_run.font.name = "Calibri"
                            value_run = sub_bullet_para.add_run(value)
                            value_run.font.name = "Calibri"

            # Double-space after each role
            if r_idx < len(roles) - 1:
                doc.add_paragraph()

        # Blank line between companies
        if c_idx < len(companies) - 1:
            doc.add_paragraph()

    # Add a line break after the last bullet before Education
    doc.add_paragraph()

    # ========== 3) EDUCATION ==========
    education_data = yaml_data["resumeStructure"].get("education", {})
    education_header = education_data.get("header", "EDUCATION")

    ed_header_para = doc.add_paragraph()
    ed_header_para.paragraph_format.space_after = Pt(0)  # Reduce space after the header
    ed_run = ed_header_para.add_run(education_header.upper())
    ed_run.font.name = "Calibri"
    ed_run.font.size = Pt(12)
    ed_run.bold = True
    ed_header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_horizontal_line(doc)

    university_name_line = education_data.get("universityNameLine", "")
    degree_info_line = education_data.get("degreeInfoLine", "")

    univ_para = doc.add_paragraph()
    univ_run = univ_para.add_run(university_name_line)
    univ_run.font.name = "Calibri"
    univ_run.font.size = Pt(12)
    univ_run.bold = True

    parts = [p.strip() for p in degree_info_line.split("|")]
    deg = parts[0] if len(parts) > 0 else ""
    location = parts[1] if len(parts) > 1 else ""

    add_left_right_paragraph(
        doc,
        left_text=deg,
        right_text=location,
        left_italic=True,
        right_italic=True,
        font_size=12
    )

    # Line break after the degree info before the next header
    doc.add_paragraph()

    # ========== 4) CERTIFICATIONS & SKILLS ==========
    csi_data = yaml_data["resumeStructure"].get("certificationsSkills", {})
    csi_header = csi_data.get("header", "CERTIFICATIONS & SKILLS")  # Updated header

    csi_header_para = doc.add_paragraph()
    csi_header_para.paragraph_format.space_after = Pt(0)  # Reduce space after the header
    csi_run = csi_header_para.add_run(csi_header.upper())
    csi_run.font.name = "Calibri"
    csi_run.font.size = Pt(12)
    csi_run.bold = True
    csi_header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_horizontal_line(doc)

    csi_bullets = csi_data.get("bullets", [])
    for bullet_obj in csi_bullets:
        for key, val in bullet_obj.items():
            if key.lower() != "interests" and val.strip():  # Exclude "Interests"
                para = doc.add_paragraph(style='List Bullet')
                label_run = para.add_run(f"{key.title()}: ")
                label_run.font.name = "Calibri"
                label_run.font.size = Pt(12)
                label_run.bold = True

                info_run = para.add_run(val)
                info_run.font.name = "Calibri"
                info_run.font.size = Pt(12)


def convert_yaml_to_word(input_dir, output_dir):
    # Prompt user for company name
    company_name = input("Enter the company name for the resume: ").strip()

    if not company_name:
        raise ValueError("Company name cannot be empty!")

    os.makedirs(output_dir, exist_ok=True)

    for filename in os.listdir(input_dir):
        if filename.endswith('.yaml'):   # ...changed extension condition...
            input_path = os.path.join(input_dir, filename)
            # Format today's date
            today_date = datetime.now().strftime("%Y-%m-%d")
            # Generate the output filename
            output_filename = f"Andrew Stasi Resume - {company_name} - {today_date}.docx"
            output_path = os.path.join(output_dir, output_filename)

            try:
                with open(input_path, 'r', encoding='utf-8') as file:
                    yaml_data = yaml.safe_load(file)   # ...changed file loading...
            except yaml.YAMLError:
                warnings.warn(
                    f"File {filename} is not a valid YAML file and will be skipped.",
                    UserWarning
                )
                continue

            doc = Document()

            # Set margins
            section = doc.sections[0]
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.gutter = 0

            # Build the resume
            create_resume_from_yaml(yaml_data, doc)

            # Save
            doc.save(output_path)
            print(f"Saved {output_filename} to {output_dir}")


if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.abspath(__file__))
    input_directory = os.path.join(base_dir, "rawjson-resume")
    output_directory = os.path.join(base_dir, "formatted-output")
    convert_yaml_to_word(input_directory, output_directory)
